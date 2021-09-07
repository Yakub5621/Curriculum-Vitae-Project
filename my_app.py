from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()


# profile picture
document.add_picture("profile.jpg", width=Inches(2.0))

# name, phone number and email details
name = input("What is your name?: ")
speak("hello " + name + " how are you today?" )

speak("What is your phone number?" )
phone_number = input("What is your phone number?: ")
email = input("What is your email?: ")

document.add_paragraph(name + " | " + phone_number +" | " + email)

# about me
document.add_heading("About me")
about_me = input("Tell me about yourself?: ")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work Experience")
paragraph = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From Date" )
to_date = input("To Date ")

paragraph.add_run(company + " ").bold = True
paragraph.add_run(from_date + "-" + to_date + "\n").italic = True
experience_details = input("Describe your experience at " + company)
paragraph.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input("Do you have more experiences? Yes or No: ")

    if has_more_experiences.lower() == "yes":
        paragraph = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From Date" )
        to_date = input("To Date ")

        paragraph.add_run(company + " ").bold = True
        paragraph.add_run(from_date + "-" + to_date + "\n").italic = True
        experience_details = input("Describe your experience at " + company )
        paragraph.add_run(experience_details)
    else:
        break

#skills
document.add_heading("Skills")
skills = input("Tell me a skill you have?: ")
paragraph = document.add_paragraph(skills)
paragraph.style = "List Bullet"

#more skills
while True:
    has_more_skills = input("Do you have more skills? Yes or No?: ")

    if has_more_skills.lower() == "yes":
        another_skill = input("Tell me another skill you have?: ")
        paragraph = document.add_paragraph(another_skill)
        paragraph.style = "List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "CV generated using Python"



document.save("cv.docx")